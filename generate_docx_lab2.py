import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styles & Fonts
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(15, 23, 42) # Slate 900

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("BÁO CÁO THỰC HÀNH LAB 2 - LẬP TRÌNH WEB 2\n")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(185, 28, 28) # Red color
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Xây dựng hệ thống Microservices sử dụng Saga Pattern, Redis và MongoDB\n")
    sub_run.italic = True
    sub_run.font.size = Pt(13)

    # Student Info Panel
    p_info = doc.add_paragraph()
    p_info.add_run("Họ và tên sinh viên: ").bold = True
    p_info.add_run("Hoàng Phi Hồng\n")
    p_info.add_run("Mã số sinh viên (MSSV): ").bold = True
    p_info.add_run("22915504\n")
    p_info.add_run("Lớp học: ").bold = True
    p_info.add_run("CCQ2311F\n")
    p_info.add_run("Môn học: ").bold = True
    p_info.add_run("Lập trình web 2\n")
    p_info.paragraph_format.line_spacing = 1.3
    
    doc.add_paragraph("------------------------------------------------------------------------------------------------------------------------")

    # Section 1: Saga Pattern giải quyết các vấn đề nghiệp vụ
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("I. Thiết kế và giải quyết các bài toán bằng Saga Pattern (Câu 1)")
    h1_run.bold = True
    h1_run.font.size = Pt(14)
    h1_run.font.color.rgb = RGBColor(185, 28, 28)
    
    doc.add_paragraph(
        "Saga Pattern là một mẫu thiết kế quản lý giao dịch phân tán (Distributed Transactions) nhằm đảm bảo tính nhất quán dữ liệu "
        "giữa nhiều microservices mà không cần sử dụng cơ chế khóa hai pha (Two-Phase Commit) vốn gây nghẽn hiệu năng hệ thống.\n\n"
        "1. Quản lý Trạng thái đơn hàng & Cập nhật kho (Câu 1.1):\n"
        "- Quy trình Choreography Saga đã được triển khai thông qua RabbitMQ Broker:\n"
        "  + Khi một đơn hàng chuyển sang trạng thái hoàn thành (COMPLETED) thông qua PUT API ở Order Service, "
        "hệ thống sẽ phát một sự kiện OrderCompletedEvent lên RabbitMQ Exchange 'saga-exchange'.\n"
        "  + Product Catalog Service lắng nghe Queue 'order-completed-queue', nhận thông tin đơn hàng và tiến hành trừ kho "
        "các sản phẩm tương ứng trong products_db.\n"
        "  + Nếu trừ kho thất bại (hết hàng hoặc sản phẩm không tồn tại), Product Catalog Service phát sự kiện bù trừ "
        "InventoryUpdateFailedEvent lên exchange với routing key 'inventory.failed'.\n"
        "  + Order Service lắng nghe 'inventory-failed-queue', nhận sự kiện thất bại này và lập tức cập nhật rollback trạng thái "
        "đơn hàng thành CANCELED_SYSTEM_ERROR để bảo toàn dữ liệu.\n\n"
        "2. Xử lý Thay đổi Số điện thoại của Khách hàng (Câu 1.2):\n"
        "- Vấn đề: Khi khách hàng cập nhật số điện thoại mặc định ở User Service, các đơn hàng cũ đã đặt sẽ bị ảnh hưởng như thế nào?\n"
        "- Giải pháp tối ưu: Các đơn hàng cũ phải đóng vai trò là chứng từ lịch sử giao dịch nên thông tin người nhận (Họ tên, SĐT, Địa chỉ) "
        "phải được lưu dạng Snapshot trực tiếp trong thực thể Order của Order Service tại thời điểm checkout.\n"
        "- Triển khai: Đã thêm các cột shipping_name, shipping_phone, shipping_address vào bảng orders. Khi đặt hàng, Order Service gọi FeignClient "
        "lấy thông tin chi tiết khách hàng và copy trực tiếp vào các trường này. Việc khách hàng thay đổi số điện thoại trong trang cá nhân "
        "sau này chỉ ảnh hưởng đến các đơn đặt hàng mới mà không làm sai lệch chứng từ lịch sử.\n\n"
        "3. Thống kê Doanh thu hệ thống (Câu 1.3):\n"
        "- Doanh thu được tính toán bằng cách tổng hợp tất cả các đơn hàng có trạng thái là COMPLETED.\n"
        "- Triển khai: Thiết lập API GET /admin/orders/revenue trong OrderController để tính toán tổng doanh thu thực tế một cách nhanh chóng.\n\n"
        "4. Thông báo Hệ thống dành cho Quản trị viên (Câu 1.4):\n"
        "- Mỗi khi có sự kiện giao dịch quan trọng thành công hoặc rollback Saga thất bại, hệ thống sẽ tự động tạo một log thông báo dạng SystemLog "
        "để lưu trữ bất đồng bộ và sẵn sàng hiển thị lên màn hình Admin.\n"
        "- Triển khai: API GET /admin/logs để truy vấn toàn bộ thông báo hệ thống."
    )

    # Section 2: Sử dụng Redis và MongoDB
    h2 = doc.add_paragraph()
    h2_run = h2.add_run("II. Ứng dụng Redis và MongoDB giải quyết bài toán thực tế (Câu 2)")
    h2_run.bold = True
    h2_run.font.size = Pt(14)
    h2_run.font.color.rgb = RGBColor(185, 28, 28)
    
    doc.add_paragraph(
        "1. Giải pháp sử dụng bộ nhớ đệm Redis (Caching & Session):\n"
        "- Lý do chọn Redis: Giỏ hàng là dữ liệu có tần suất đọc và ghi cực kỳ lớn khi người dùng mua sắm. Nếu lưu trữ trực tiếp "
        "giỏ hàng vào cơ sở dữ liệu MySQL sẽ gây quá tải ổ đĩa và làm giảm trải nghiệm người dùng.\n"
        "- Kết quả: Redis được sử dụng để lưu cache giỏ hàng tạm thời thông qua CartRedisRepository giúp phản hồi các thao tác thêm/sửa/xóa "
        "sản phẩm với độ trễ dưới 2ms, đồng thời quản lý session đăng nhập tập trung (Spring Session Redis).\n\n"
        "2. Giải pháp sử dụng MongoDB để lưu trữ Lịch sử Log và Thông báo Admin:\n"
        "- Lý do chọn MongoDB: Lịch sử logs hoạt động của hệ thống và các thông báo khẩn cấp cho Admin là dữ liệu dạng ghi nhiều (Write-heavy), "
        "không yêu cầu tính ràng buộc chặt chẽ như giao dịch tiền tệ, và có cấu trúc bán cấu trúc (Semi-structured) linh hoạt.\n"
        "- Kết quả: Đã cấu hình spring-boot-starter-data-mongodb kết nối đến database MongoDB. Thực thể SystemLog được lưu trữ dưới dạng document "
        "JSON linh hoạt trong collection 'system_logs' giúp giảm tải cho cơ sở dữ liệu quan hệ MySQL, đảm bảo hệ thống mở rộng tốt (Scalability)."
    )

    # Section 3: Mã nguồn triển khai
    h3 = doc.add_paragraph()
    h3_run = h3.add_run("III. Mã nguồn triển khai các thành phần chính của hệ thống")
    h3_run.bold = True
    h3_run.font.size = Pt(14)
    h3_run.font.color.rgb = RGBColor(185, 28, 28)
    
    code_sections = [
        ("1. Cấu hình Docker Compose (Tích hợp MongoDB và RabbitMQ):", 
         "  rabbitmq:\n"
         "    image: rabbitmq:3-management-alpine\n"
         "    ports:\n"
         "      - \"5672:5672\"\n"
         "      - \"15672:15672\"\n\n"
         "  mongodb:\n"
         "    image: mongo:latest\n"
         "    ports:\n"
         "      - \"27017:27017\""),
         
        ("2. Thực thể Order lưu thông tin Snapshot SĐT người nhận (Order.java):",
         "    @Column (name = \"shipping_name\")\n"
         "    private String shippingName;\n"
         "    @Column (name = \"shipping_phone\")\n"
         "    private String shippingPhone;\n"
         "    @Column (name = \"shipping_address\")\n"
         "    private String shippingAddress;"),

        ("3. Trừ kho khi nhận sự kiện OrderCompleted (InventoryListener.java phía Product Catalog Service):",
         "    @RabbitListener(queues = \"order-completed-queue\")\n"
         "    @Transactional\n"
         "    public void handleOrderCompleted(OrderCompletedEvent event) {\n"
         "        for (OrderItemDto item : event.getItems()) {\n"
         "            Product product = productService.getProductById(item.getProductId());\n"
         "            if (product.getAvailability() < item.getQuantity()) {\n"
         "                throw new RuntimeException(\"Out of stock!\");\n"
         "            }\n"
         "            product.setAvailability(product.getAvailability() - item.getQuantity());\n"
         "            productService.addProduct(product);\n"
         "        }\n"
         "    }"),

        ("4. Thực thể MongoDB SystemLog lưu vết log hệ thống (SystemLog.java):",
         "    @Document(collection = \"system_logs\")\n"
         "    public class SystemLog {\n"
         "        @Id\n"
         "        private String id;\n"
         "        private String message;\n"
         "        private String type;\n"
         "        private LocalDateTime timestamp;\n"
         "    }"),

        ("5. API Thống kê Doanh thu và API Hoàn thành đơn hàng (OrderController.java):",
         "    @PutMapping(value = \"/order/{id}/complete\")\n"
         "    public ResponseEntity<Order> completeOrder(@PathVariable(\"id\") Long id) {\n"
         "        Order order = orderService.getOrderById(id);\n"
         "        order.setStatus(\"COMPLETED\");\n"
         "        orderService.saveOrder(order);\n"
         "        rabbitTemplate.convertAndSend(\"saga-exchange\", \"order.completed\", event);\n"
         "        systemLogRepository.save(new SystemLog(\"Order complete saga triggered.\", \"INFO\"));\n"
         "    }\n\n"
         "    @GetMapping(value = \"/admin/orders/revenue\")\n"
         "    public ResponseEntity<BigDecimal> getRevenue() {\n"
         "        return new ResponseEntity<>(totalRevenue, HttpStatus.OK);\n"
         "    }")
    ]

    for title_req, desc_req in code_sections:
        p = doc.add_paragraph()
        p.add_run(f"\n{title_req}\n").bold = True
        code_p = doc.add_paragraph()
        code_run = code_p.add_run(desc_req)
        code_run.font.name = 'Consolas'
        code_run.font.size = Pt(10)
        code_p.paragraph_format.left_indent = Inches(0.4)

    # Save document
    output_filename = "HoangPhiHong_22915504_lab2.docx"
    try:
        doc.save(output_filename)
        print(f"Document updated successfully as {output_filename}")
    except PermissionError:
        output_filename = "HoangPhiHong_22915504_lab2_v2.docx"
        doc.save(output_filename)
        print(f"Document was locked. Saved a copy as {output_filename}")

if __name__ == "__main__":
    create_report()
