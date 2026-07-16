import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styles & Fonts
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(15, 23, 42) # Slate 900

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("BÁO CÁO THỰC HÀNH LAB 1 - LẬP TRÌNH WEB 2\n")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(185, 28, 28) # Red color from image accents
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Phân tích Monolithic, thiết kế Microservices bằng Spring Boot, OpenFeign và RabbitMQ\n")
    sub_run.italic = True
    sub_run.font.size = Pt(13)

    # Student Info Panel
    p_info = doc.add_paragraph()
    p_info.add_run("Họ và tên sinh viên: ").bold = True
    p_info.add_run("Hoàng Phi Hồng\n")
    p_info.add_run("Mã số sinh viên (MSSV): ").bold = True
    p_info.add_run("22915504\n")
    p_info.add_run("Lớp học: ").bold = True
    p_info.add_run("CCQ2311F\n")
    p_info.add_run("Môn học: ").bold = True
    p_info.add_run("Lập trình web 2\n")
    p_info.paragraph_format.line_spacing = 1.3
    
    doc.add_paragraph("------------------------------------------------------------------------------------------------------------------------")

    # Section 1: Monolithic vs Microservices
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("I. Phân tích hệ thống Monolithic và quá trình tách thành Microservices")
    h1_run.bold = True
    h1_run.font.size = Pt(14)
    h1_run.font.color.rgb = RGBColor(185, 28, 28)
    
    doc.add_paragraph(
        "1. Hệ thống Monolithic quản lý bán hàng đơn giản (gồm Product và Order):\n"
        "- Trong kiến trúc Monolithic cũ, module quản lý Sản phẩm (Product) và module Quản lý Đơn hàng (Order) được triển khai chung "
        "trong một ứng dụng (Single Deployable Unit) và chia sẻ cùng một cơ sở dữ liệu duy nhất.\n"
        "- Điểm yếu: Khi số lượng truy vấn tăng cao, CSDL duy nhất dễ bị nghẽn (Bottleneck). Nếu module Order bị quá tải hoặc lỗi, "
        "toàn bộ hệ thống (bao gồm cả xem sản phẩm) cũng sẽ bị ngừng hoạt động.\n\n"
        "2. Phân tách thành kiến trúc Microservices độc lập:\n"
        "- Hệ thống đã được phân tách thành 3 service chính độc lập có liên kết database riêng:\n"
        "  + Product Service: Quản lý thông tin chi tiết sản phẩm, danh mục và cập nhật tồn kho (Database: products_db).\n"
        "  + User Service: Quản lý thông tin khách hàng, tài khoản đăng nhập và phân quyền (Database: users_db).\n"
        "  + Order Service: Xử lý giỏ hàng, tạo đơn đặt hàng và tính toán tổng số tiền (Database: orders_db).\n"
        "- Các dịch vụ độc lập này tự quản lý CSDL của riêng mình, giao tiếp với nhau qua API HTTP (đồng bộ) hoặc Message Broker (bất đồng bộ)."
    )

    # Section 2: Triển khai các yêu cầu kỹ thuật 1 đến 6
    h2 = doc.add_paragraph()
    h2_run = h2.add_run("II. Kết quả triển khai các yêu cầu kỹ thuật của bài Lab")
    h2_run.bold = True
    h2_run.font.size = Pt(14)
    h2_run.font.color.rgb = RGBColor(185, 28, 28)
    
    requirements = [
        ("1. Sử dụng OpenFeign trong dự án:", 
         "Dự án sử dụng Spring Cloud OpenFeign để Order Service giao tiếp đồng bộ với Product Service và User Service. "
         "Chúng tôi định nghĩa các Client Interface (`ProductClient`, `UserClient`) kèm theo các Request Mapping của Spring Web, "
         "giúp loại bỏ hoàn toàn mã nguồn viết tay HttpClient rườm rà."),
         
        ("2. Gọi API để lấy thông tin khách hàng và thông tin sản phẩm:", 
         "Order Service gọi endpoint `/accounts/users/{id}` (phía User Service) để xác thực thông tin khách hàng "
         "và gọi endpoint `/catalog/products/{id}` (phía Product Service) để lấy chi tiết sản phẩm trước khi tạo đơn hàng."),
         
        ("3. Tra cứu thông tin sản phẩm khi tạo đơn hàng:", 
         "Khi thực hiện tạo đơn hàng, Order Service tự động gọi API lấy thông tin chi tiết của từng sản phẩm trong giỏ hàng "
         "nhằm mục đích kiểm tra giá bán thực tế và số lượng khả dụng trong kho hàng của Product Service."),
         
        ("4. Kiểm tra thông tin khách hàng trước khi đặt hàng:", 
         "Nhằm tránh tình trạng tạo đơn hàng ảo hoặc cho người dùng không tồn tại, Order Service sẽ gửi request kiểm tra trạng thái "
         "tài khoản khách hàng qua User Service. Nếu tài khoản không hợp lệ hoặc bị khóa, quy trình đặt hàng sẽ bị từ chối."),
         
        ("5. Kết quả trả về thể hiện đầy đủ thông tin (Đặt nhiều sản phẩm):", 
         "Khi tạo đơn hàng thành công, kết quả trả về phía client chứa đầy đủ thông tin bao gồm: Thông tin chi tiết của khách hàng "
         "(Họ tên, email, SĐT), danh sách tất cả sản phẩm đặt mua kèm theo đơn giá, số lượng đặt mua tương ứng của từng sản phẩm, "
         "và tổng giá trị thanh toán của đơn hàng (Total Amount)."),
         
        ("6. Xử lý bất đồng bộ qua RabbitMQ sau khi đặt hàng:", 
         "Sau khi lưu đơn hàng vào database thành công, Order Service gửi sự kiện `OrderCreatedEvent` lên RabbitMQ Broker. "
         "Notification Service lắng nghe hàng đợi (Queue) sẽ nhận sự kiện này bất đồng bộ và giả lập gửi email xác nhận đặt hàng cho "
         "khách hàng. Quy trình này không gây chậm trễ thời gian phản hồi cho khách hàng trên trình duyệt.")
    ]
    
    for title_req, desc_req in requirements:
        p = doc.add_paragraph()
        p.add_run(f"- {title_req} ").bold = True
        p.add_run(desc_req)
        p.paragraph_format.left_indent = Inches(0.2)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Section 3: Giải thích lựa chọn công nghệ
    h3 = doc.add_paragraph()
    h3_run = h3.add_run("III. Tìm hiểu và giải thích lựa chọn công nghệ (Yêu cầu 1)")
    h3_run.bold = True
    h3_run.font.size = Pt(14)
    h3_run.font.color.rgb = RGBColor(185, 28, 28)
    
    # OpenFeign vs RestTemplate
    p_feign = doc.add_paragraph()
    p_feign.add_run("1. Tại sao chọn OpenFeign thay vì RestTemplate?\n").bold = True
    p_feign.add_run(
        "- Khai báo (Declarative): OpenFeign giúp định nghĩa HTTP Client chỉ bằng một Interface và các Annotation của Spring MVC. "
        "Ngược lại, RestTemplate đòi hỏi viết code mệnh lệnh (Imperative), tự xử lý URI, HttpEntity, parse JSON và handle lỗi thủ công.\n"
        "- Tự động tích hợp Service Discovery: OpenFeign liên kết trực tiếp với Eureka Server, tự động cân bằng tải (LoadBalancer) "
        "khi gọi dịch vụ chỉ bằng tên định danh (như `@FeignClient(name = \"user-service\")`), rất dễ bảo trì.\n"
        "- Trường hợp phù hợp:\n"
        "  + OpenFeign phù hợp nhất khi giao tiếp giữa các dịch vụ nội bộ (Internal communication) trong cụm microservices.\n"
        "  + RestTemplate (hoặc WebClient) phù hợp khi gọi các API của bên thứ ba bên ngoài hệ thống (External payment gateways, SMS OTP APIs) "
        "do không cần quản lý qua Service Registry."
    )
    p_feign.paragraph_format.left_indent = Inches(0.2)
    
    # RabbitMQ vs Kafka
    p_mq = doc.add_paragraph()
    p_mq.add_run("2. Tại sao chọn RabbitMQ thay vì Kafka?\n").bold = True
    p_mq.add_run(
        "- Mô hình Broker: RabbitMQ là Message Broker truyền thống dựa trên RAM (xóa tin nhắn ngay khi Consumer xác nhận tiêu thụ thành công), "
        "phù hợp cho xử lý tác vụ nền ngắt quãng. Kafka là Commit Log phân tán ghi tuần tự trên đĩa cứng, tối ưu cho luồng dữ liệu liên tục.\n"
        "- Độ phức tạp: RabbitMQ dễ cấu hình định tuyến thông qua Exchanges/Queues linh hoạt. Kafka phức tạp hơn, yêu cầu quản lý ZooKeeper/KRaft "
        "và phân vùng (Partitions).\n"
        "- Trường hợp phù hợp:\n"
        "  + RabbitMQ phù hợp cho các nghiệp vụ gửi thông báo, xử lý hàng đợi giao dịch, tác vụ bất đồng bộ ngầm đơn giản (như gửi email đặt hàng).\n"
        "  + Kafka phù hợp cho thu thập log tập trung (Log Aggregation), Clickstream phân tích hành vi người dùng thời gian thực, IoT Data Streaming."
    )
    p_mq.paragraph_format.left_indent = Inches(0.2)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Section 4: Trình bày kết quả
    h4 = doc.add_paragraph()
    h4_run = h4.add_run("IV. Trình bày kết quả vận hành hệ thống (Yêu cầu 2)")
    h4_run.bold = True
    h4_run.font.size = Pt(14)
    h4_run.font.color.rgb = RGBColor(185, 28, 28)
    
    doc.add_paragraph(
        "Dưới đây là hình ảnh thực tế vận hành hệ thống Microservices bao gồm giao diện chính, bảng điều khiển quản trị, "
        "và hoạt động của các form CRUD đã thiết lập:"
    )

    # Embed Images
    brain_dir = r"C:\Users\Phi Hong\Downloads\e-commerce-microservices-master"
    # Copy images to workspace if not present, but wait, they are in brain_dir or we can just access them from brain_dir
    # Let's specify exact paths
    base_brain = r"C:\Users\Phi Hong\.gemini\antigravity\brain\70bacb66-0e80-4c51-b54d-4190039ae9ab"
    
    images_to_add = [
        ("browser_homepage_1781513173611.png", "Hình 1: Giao diện Trang chủ (Storefront) phía khách hàng hiển thị đầy đủ sản phẩm."),
        ("admin_login_page_1781512231861.png", "Hình 2: Giao diện Đăng nhập hệ thống Quản trị (Admin Login)."),
        ("admin_dashboard_1781512255030.png", "Hình 3: Bảng điều khiển quản trị Admin Dashboard hiển thị thống kê thời gian thực."),
        ("admin_products_list_1781512269434.png", "Hình 4: Danh sách quản lý sản phẩm trong hệ thống Admin CMS."),
        ("admin_product_modal_1781512282522.png", "Hình 5: Form cập nhật thông tin sản phẩm (Edit Product Modal) dạng Glassmorphism."),
        ("admin_users_list_1781512304229.png", "Hình 6: Danh sách quản lý tài khoản người dùng và phân quyền hệ thống.")
    ]
    
    for filename, caption in images_to_add:
        img_path = os.path.join(base_brain, filename)
        if os.path.exists(img_path):
            try:
                doc.add_paragraph().paragraph_format.space_before = Pt(6)
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.add_run().add_picture(img_path, width=Inches(5.5))
                
                p_cap = doc.add_paragraph()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_run = p_cap.add_run(caption)
                cap_run.italic = True
                cap_run.font.size = Pt(10)
            except Exception as e:
                print(f"Error adding image {filename}: {e}")
        else:
            p_missing = doc.add_paragraph()
            p_missing.add_run(f"[Hình ảnh minh họa: {filename} không tìm thấy]").italic = True

    # Save document
    output_filename = "HoangPhiHong_22915504_lab1.docx"
    try:
        doc.save(output_filename)
        print(f"Document updated successfully as {output_filename}")
    except PermissionError:
        output_filename = "HoangPhiHong_22915504_lab1_v2.docx"
        doc.save(output_filename)
        print(f"Document was locked. Saved a copy as {output_filename}")

if __name__ == "__main__":
    create_report()
